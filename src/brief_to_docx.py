"""
Convert executive briefs from Markdown to DOCX format.

Handles markdown formatting (headings, bold, italics, lists, links, etc.)
and produces professional Word documents suitable for distribution.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _parse_markdown_lines(text: str) -> List[Tuple[str, str]]:
    """
    Parse markdown into (type, content) tuples.
    Types: 'heading1', 'heading2', 'heading3', 'paragraph', 'list_item', 'blockquote', 'code_block'
    """
    lines = text.split('\n')
    parsed: List[Tuple[str, str]] = []
    in_code_block = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block (triple backticks)
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            parsed.append(('code_block', line))
            i += 1
            continue

        # Skip empty lines but preserve them for spacing
        if not line.strip():
            parsed.append(('blank', ''))
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            parsed.append(('heading1', line[2:].strip()))
        elif line.startswith('## '):
            parsed.append(('heading2', line[3:].strip()))
        elif line.startswith('### '):
            parsed.append(('heading3', line[4:].strip()))
        # List items
        elif line.strip().startswith('- '):
            parsed.append(('list_item', line.strip()[2:].strip()))
        # Blockquotes
        elif line.strip().startswith('> '):
            parsed.append(('blockquote', line.strip()[2:].strip()))
        # Details/collapsible sections
        elif line.strip().lower() in ('<details>', '</details>', '<summary>', '</summary>'):
            # Skip HTML details tags
            continue
        # Regular paragraph
        else:
            parsed.append(('paragraph', line.strip()))

        i += 1

    return parsed


def _format_inline(text: str) -> List[Tuple[str, str, dict]]:
    """
    Parse inline formatting (bold, italics, links, code) into runs.
    Returns list of (text, format_type, style_dict) tuples.
    """
    runs: List[Tuple[str, str, dict]] = []

    # Pattern to find formatting: **bold**, *italic*, `code`, [link](url)
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\)|[^*`\[]+)'

    for match in re.finditer(pattern, text):
        token = match.group(0)

        if token.startswith('**') and token.endswith('**'):
            # Bold
            runs.append((token[2:-2], 'text', {'bold': True}))
        elif token.startswith('*') and token.endswith('*'):
            # Italic
            runs.append((token[1:-1], 'text', {'italic': True}))
        elif token.startswith('`') and token.endswith('`'):
            # Code
            runs.append((token[1:-1], 'code', {}))
        elif token.startswith('[') and '](' in token:
            # Link
            m = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', token)
            if m:
                link_text, url = m.groups()
                runs.append((link_text, 'link', {'url': url}))
            else:
                runs.append((token, 'text', {}))
        else:
            # Regular text
            if token:
                runs.append((token, 'text', {}))

    return runs if runs else [(text, 'text', {})]


def _add_paragraph_with_formatting(doc: Document, text: str, style: str = 'Normal') -> None:
    """Add a paragraph with inline formatting (bold, italic, links, etc.)"""
    p = doc.add_paragraph(style=style)

    if not text.strip():
        return

    runs_data = _format_inline(text)
    for run_text, format_type, style_dict in runs_data:
        run = p.add_run(run_text)

        if format_type == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(50, 50, 50)
        elif format_type == 'link':
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
            # Store URL in style_dict for potential hyperlink handling
        elif 'bold' in style_dict and style_dict['bold']:
            run.font.bold = True
        elif 'italic' in style_dict and style_dict['italic']:
            run.font.italic = True

    # Adjust paragraph alignment
    if 'center' in style.lower():
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def markdown_to_docx(markdown_text: str, title: str = "Executive Brief") -> BytesIO:
    """
    Convert markdown text to a Word document (.docx).

    Args:
        markdown_text: The markdown content of the brief
        title: Document title (appears in title section)

    Returns:
        BytesIO object containing the .docx file data
    """
    doc = Document()

    # Add title
    title_para = doc.add_paragraph(title, style='Heading 1')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a blank line
    doc.add_paragraph()

    # Parse and add content
    parsed_lines = _parse_markdown_lines(markdown_text)

    code_block_buffer: List[str] = []
    in_code_block = False

    for line_type, content in parsed_lines:
        if line_type == 'heading1':
            doc.add_paragraph(content, style='Heading 1')
        elif line_type == 'heading2':
            doc.add_paragraph(content, style='Heading 2')
        elif line_type == 'heading3':
            doc.add_paragraph(content, style='Heading 3')
        elif line_type == 'list_item':
            p = doc.add_paragraph(content, style='List Bullet')
            # Format inline content in list items
            _format_list_item_inline(p, content)
        elif line_type == 'blockquote':
            p = doc.add_paragraph(content, style='Quote')
            # Add left border to blockquote
            _add_blockquote_formatting(p)
        elif line_type == 'code_block':
            code_block_buffer.append(content)
            if not content.strip():  # End of code block
                if code_block_buffer:
                    code_text = '\n'.join(code_block_buffer[:-1])
                    _add_code_block(doc, code_text)
                    code_block_buffer = []
        elif line_type == 'blank':
            # Only add blank lines if they won't create excessive spacing
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                pass  # Implicit spacing
        elif line_type == 'paragraph':
            _add_paragraph_with_formatting(doc, content)

    # Handle any remaining code block
    if code_block_buffer:
        code_text = '\n'.join(code_block_buffer)
        _add_code_block(doc, code_text)

    # Convert to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _format_list_item_inline(paragraph, text: str) -> None:
    """Format inline markdown in a list item paragraph."""
    # Clear the default run
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    runs_data = _format_inline(text)
    for run_text, format_type, style_dict in runs_data:
        run = paragraph.add_run(run_text)

        if format_type == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(50, 50, 50)
        elif format_type == 'link':
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
        elif 'bold' in style_dict and style_dict['bold']:
            run.font.bold = True
        elif 'italic' in style_dict and style_dict['italic']:
            run.font.italic = True


def _add_blockquote_formatting(paragraph) -> None:
    """Add visual styling to a blockquote paragraph."""
    # Add left border
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')  # Size in eighths of a point
    left.set(qn('w:space'), '24')
    left.set(qn('w:color'), 'CCCCCC')

    pBdr.append(left)
    pPr.append(pBdr)

    # Add gray background
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

    # Indent left
    ind = paragraph._element.pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '720')  # 0.5 inch


def _add_code_block(doc: Document, code_text: str) -> None:
    """Add a code block with monospace formatting."""
    p = doc.add_paragraph(style='Normal')

    # Add gray background
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)

    # Add padding
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:right'), '360')
    p._element.pPr.append(ind)

    # Add border
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    p._element.pPr.append(pBdr)

    # Add code text
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)
