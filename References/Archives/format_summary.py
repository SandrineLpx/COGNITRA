#!/usr/bin/env python3
"""
format_summary.py - Format and structure automotive market intelligence summaries

This script takes raw summary text and structures it with proper metadata
for SharePoint integration and consistent formatting.

Usage:
    python format_summary.py input.txt [--output output.txt] [--format docx|txt]
    python format_summary.py --interactive

Requirements:
    - python-docx (for Word document generation)
    - json (for metadata extraction)
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple


import json
from pathlib import Path

def load_taxonomy(path: str = "taxonomy.json") -> dict:
    """
    Load canonical topics and region rules from taxonomy.json.
    """
    return json.loads(Path(path).read_text(encoding="utf-8"))

# Country extraction uses pycountry (installed) and simple word-boundary matching.
try:
    import pycountry
except Exception:
    pycountry = None

def _build_country_patterns():
    patterns = {}
    if pycountry is None:
        return patterns
    names = set()
    for c in list(pycountry.countries):
        for attr in ("name", "official_name", "common_name"):
            v = getattr(c, attr, None)
            if v:
                names.add(v)
        # Include alpha_2/alpha_3? Not helpful in prose, skip.
    # Add common variants
    names.update(["UK", "U.K.", "United States", "U.S.", "USA", "U.S.A.", "Russia", "Czech Republic", "Czechia", "South Korea", "North Korea"])
    # Sort longer first to avoid partial matches
    for name in sorted(names, key=len, reverse=True):
        # escape dots etc; word boundaries with spaces/hyphens handled loosely
        pat = re.compile(r'(?<![A-Za-z])' + re.escape(name) + r'(?![A-Za-z])', re.IGNORECASE)
        patterns[name] = pat
    return patterns

_COUNTRY_PATTERNS = _build_country_patterns()

def extract_country_mentions(text: str) -> list:
    if not text:
        return []
    found = set()
    for name, pat in _COUNTRY_PATTERNS.items():
        if pat.search(text):
            # normalize some aliases
            key = name
            if key.lower() in {"u.s.", "u.s.a.", "usa", "united states"}:
                key = "United States"
            if key.lower() in {"uk", "u.k.", "united kingdom", "britain", "great britain"}:
                key = "United Kingdom"
            found.add(key)
    return sorted(found)

def rollup_regions(text: str, country_mentions: list) -> dict:
    """
    Returns:
      regions_mentioned, regions_relevant_to_kiekert
    """
    taxonomy = load_taxonomy()
    footprint = set(taxonomy.get("footprint_regions", []))

    t = (text or "").lower()
    regions_mentioned = set()

    # Broad mentions -> Europe rollup
    europe_terms = {"europe", "eu", "european", "emea", "european union"}
    if any(term in t for term in europe_terms):
        regions_mentioned.add("Europe (including Russia)")

    # Broad mentions -> Africa rollup
    if "africa" in t or "african" in t:
        regions_mentioned.add("Africa")

    # Country-based rollups
    europe_countries = {
        "United Kingdom","Turkey","Russia","France","Germany","Spain","Italy","Czech Republic","Czechia",
        "Poland","Netherlands","Belgium","Sweden","Norway","Finland","Denmark","Austria","Switzerland",
        "Portugal","Greece","Hungary","Romania","Bulgaria","Slovakia","Slovenia","Croatia","Serbia",
        "Ukraine","Ireland","Iceland","Estonia","Latvia","Lithuania"
    }
    africa_countries = {"Morocco","South Africa"}

    for c in country_mentions or []:
        if c in {"India","China","Mexico","Thailand","United States"}:
            # map to footprint region labels
            if c == "United States":
                regions_mentioned.add("US")
            else:
                regions_mentioned.add(c)
        if c in europe_countries:
            regions_mentioned.add("Europe (including Russia)")
        if c in africa_countries:
            regions_mentioned.add("Africa")

    regions_relevant = sorted(set(regions_mentioned).intersection(footprint))
    return {
        "regions_mentioned": sorted(regions_mentioned),
        "regions_relevant_to_kiekert": regions_relevant
    }



def enrich_with_region_tags(summary: dict) -> dict:
    """
    Adds country_mentions / regions_mentioned / regions_relevant_to_kiekert to the summary dict.
    """
    text_fields = []
    for k in ("title", "summary", "key_developments", "strategic_implications", "raw_text", "content"):
        v = summary.get(k)
        if isinstance(v, str):
            text_fields.append(v)
        elif isinstance(v, list):
            text_fields.extend([x for x in v if isinstance(x, str)])
    joined = "\n".join(text_fields)
    countries = summary.get("country_mentions") or extract_country_mentions(joined)
    summary["country_mentions"] = countries
    region_info = rollup_regions(joined, countries)
    summary.setdefault("regions_mentioned", region_info["regions_mentioned"])
    summary.setdefault("regions_relevant_to_kiekert", region_info["regions_relevant_to_kiekert"])
    return summary




try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word document generation will not be available.")
    print("Install with: pip install python-docx")


class SummaryFormatter:
    """Format and structure market intelligence summaries"""
    
    TAXONOMY = load_taxonomy()
    TOPICS = TAXONOMY["topics"]
    PRIORITIES = ["High", "Medium", "Low"]
    
    SOURCES = [
        "Bloomberg",
        "Automotive News",
        "Reuters",
        "Patent",
        "Press Release",
        "Other"
    ]
    
    def __init__(self):
        self.metadata = {}
        self.content = {}
    
    def parse_raw_summary(self, text: str) -> Dict:
        """
        Parse raw summary text and extract structured components
        
        Args:
            text: Raw summary text
            
        Returns:
            Dictionary with structured components
        """
        result = {
            "title": self._extract_title(text),
            "source": self._extract_source(text),
            "date_published": self._extract_date(text),
            "url": self._extract_url(text),
            "key_developments": self._extract_key_developments(text),
            "companies": self._extract_companies(text),
            "topics": self._extract_topics(text),
            "priority": self._extract_priority(text),
            "strategic_implications": self._extract_strategic_implications(text),
            "recommended_actions": self._extract_recommended_actions(text),
        }
        
        return result
    
    def _extract_title(self, text: str) -> str:
        """Extract or generate title from text"""
        # Look for explicit title markers
        patterns = [
            r"TITLE:\s*(.+?)(?:\n|$)",
            r"Title:\s*(.+?)(?:\n|$)",
            r"^(.+?)(?:\n|$)"  # First line if no explicit title
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return match.group(1).strip()
        
        return "Untitled Summary"
    
    def _extract_source(self, text: str) -> str:
        """Extract source publication"""
        pattern = r"SOURCE:\s*([^,\n]+)"
        match = re.search(pattern, text, re.IGNORECASE)
        
        if match:
            source = match.group(1).strip()
            # Normalize to known sources
            for known_source in self.SOURCES:
                if known_source.lower() in source.lower():
                    return known_source
            return "Other"
        
        return "Unknown"
    
    def _extract_date(self, text: str) -> Optional[str]:
        """Extract publication date"""
        # Look for various date formats
        patterns = [
            r"(?:Date|Published):\s*(\w+ \d{1,2},? \d{4})",
            r"(\d{4}-\d{2}-\d{2})",
            r"(\w+ \d{1,2},? \d{4})",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_url(self, text: str) -> Optional[str]:
        """Extract original article URL"""
        pattern = r"URL:\s*(https?://[^\s\n]+)"
        match = re.search(pattern, text, re.IGNORECASE)
        
        if match:
            return match.group(1).strip()
        
        # Try to find any URL in the text
        url_pattern = r"https?://[^\s\n]+"
        match = re.search(url_pattern, text)
        if match:
            return match.group(0).strip()
        
        return None
    
    def _extract_key_developments(self, text: str) -> List[str]:
        """Extract key developments bullets"""
        # Look for KEY DEVELOPMENTS section
        pattern = r"KEY DEVELOPMENTS?:?\s*\n((?:â€¢|•|\*|-|\d+\.)\s*.+(?:\n(?:â€¢|•|\*|-|\d+\.)\s*.+)*)"
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        
        if match:
            bullets_text = match.group(1)
            bullets = re.findall(r"(?:â€¢|•|\*|-|\d+\.)\s*(.+)", bullets_text)
            return [b.strip() for b in bullets if b.strip()]
        
        return []
    
    def _extract_companies(self, text: str) -> List[str]:
        """Extract mentioned companies"""
        pattern = r"COMPANIES MENTIONED?:?\s*(.+?)(?:\n|$)"
        match = re.search(pattern, text, re.IGNORECASE)
        
        if match:
            companies = match.group(1).strip()
            # Split by commas, semicolons, or "and"
            company_list = re.split(r'[,;]|\s+and\s+', companies)
            return [c.strip() for c in company_list if c.strip()]
        
        return []
    
    def _extract_topics(self, text: str) -> List[str]:
        """Extract topic categorization"""
        pattern = r"TOPICS?:?\s*(.+?)(?:\n|$)"
        match = re.search(pattern, text, re.IGNORECASE)
        
        if match:
            topics = match.group(1).strip()
            topic_list = re.split(r'[,;]', topics)
            
            # Normalize to known topics
            normalized = []
            for topic in topic_list:
                topic = topic.strip()
                for known_topic in self.TOPICS:
                    if known_topic.lower() in topic.lower():
                        if known_topic not in normalized:
                            normalized.append(known_topic)
                        break
            
            return normalized if normalized else ["Other"]
        
        return []
    
    def _extract_priority(self, text: str) -> str:
        """Extract priority assessment"""
        pattern = r"PRIORITY:?\s*(\w+)"
        match = re.search(pattern, text, re.IGNORECASE)
        
        if match:
            priority = match.group(1).strip().capitalize()
            if priority in self.PRIORITIES:
                return priority
        
        return "Medium"  # Default
    
    def _extract_strategic_implications(self, text: str) -> str:
        """Extract strategic implications section"""
        pattern = r"STRATEGIC IMPLICATIONS?:?\s*\n(.+?)(?:\n\n|\n[A-Z]{3,}:|\Z)"
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            return match.group(1).strip()
        
        return ""
    
    def _extract_recommended_actions(self, text: str) -> List[str]:
        """Extract recommended actions"""
        pattern = r"RECOMMENDED ACTIONS?:?\s*\n((?:â€¢|•|\*|-|\d+\.)\s*.+(?:\n(?:â€¢|•|\*|-|\d+\.)\s*.+)*)"
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        
        if match:
            actions_text = match.group(1)
            actions = re.findall(r"(?:â€¢|•|\*|-|\d+\.)\s*(.+)", actions_text)
            return [a.strip() for a in actions if a.strip()]
        
        return []
    
    def format_text_output(self, data: Dict) -> str:
        """Format structured data as text summary"""
        lines = []
        
        # Title and metadata
        lines.append("=" * 70)
        lines.append(data.get("title", "Untitled Summary"))
        lines.append("=" * 70)
        lines.append("")
        
        if data.get("source"):
            source_line = f"SOURCE: {data['source']}"
            if data.get("date_published"):
                source_line += f", {data['date_published']}"
            lines.append(source_line)
        
        if data.get("url"):
            lines.append(f"URL: {data['url']}")
        
        lines.append("")
        
        # Key developments
        if data.get("key_developments"):
            lines.append("KEY DEVELOPMENTS:")
            for dev in data["key_developments"]:
                lines.append(f"• {dev}")
            lines.append("")
        
        # Metadata
        if data.get("companies"):
            lines.append(f"COMPANIES MENTIONED: {', '.join(data['companies'])}")
        
        if data.get("topics"):
            lines.append(f"TOPICS: {', '.join(data['topics'])}")
        
        if data.get("priority"):
            lines.append(f"PRIORITY: {data['priority']}")
        
        lines.append("")
        
        # Strategic implications
        if data.get("strategic_implications"):
            lines.append("STRATEGIC IMPLICATIONS:")
            lines.append(data["strategic_implications"])
            lines.append("")
        
        # Recommended actions
        if data.get("recommended_actions"):
            lines.append("RECOMMENDED ACTIONS:")
            for i, action in enumerate(data["recommended_actions"], 1):
                lines.append(f"{i}. {action}")
            lines.append("")
        
        # Footer
        lines.append("=" * 70)
        lines.append(f"Summary generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append("=" * 70)
        
        return "\n".join(lines)
    
    def format_docx_output(self, data: Dict) -> Document:
        """Format structured data as Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document generation")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(data.get("title", "Untitled Summary"), level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Metadata table
        doc.add_paragraph()  # Spacing
        
        if data.get("source"):
            p = doc.add_paragraph()
            p.add_run("Source: ").bold = True
            source_text = data["source"]
            if data.get("date_published"):
                source_text += f", {data['date_published']}"
            p.add_run(source_text)
        
        if data.get("url"):
            p = doc.add_paragraph()
            p.add_run("URL: ").bold = True
            p.add_run(data["url"])
        
        doc.add_paragraph()  # Spacing
        
        # Key developments
        if data.get("key_developments"):
            doc.add_heading("Key Developments", level=2)
            for dev in data["key_developments"]:
                doc.add_paragraph(dev, style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # Metadata section
        p = doc.add_paragraph()
        p.add_run("Companies Mentioned: ").bold = True
        p.add_run(", ".join(data.get("companies", ["None"])))
        
        p = doc.add_paragraph()
        p.add_run("Topics: ").bold = True
        p.add_run(", ".join(data.get("topics", ["Uncategorized"])))
        
        p = doc.add_paragraph()
        p.add_run("Priority: ").bold = True
        priority_run = p.add_run(data.get("priority", "Medium"))
        # Color code priority
        if data.get("priority") == "High":
            priority_run.font.color.rgb = RGBColor(255, 0, 0)
        elif data.get("priority") == "Low":
            priority_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Spacing
        
        # Strategic implications
        if data.get("strategic_implications"):
            doc.add_heading("Strategic Implications", level=2)
            doc.add_paragraph(data["strategic_implications"])
        
        # Recommended actions
        if data.get("recommended_actions"):
            doc.add_heading("Recommended Actions", level=2)
            for i, action in enumerate(data["recommended_actions"], 1):
                doc.add_paragraph(f"{i}. {action}")
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            f"Summary generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        return doc


def main():
    parser = argparse.ArgumentParser(
        description="Format automotive market intelligence summaries"
    )
    parser.add_argument(
        "input",
        nargs="?",
        help="Input file containing raw summary text"
    )
    parser.add_argument(
        "--output",
        "-o",
        help="Output file path"
    )
    parser.add_argument(
        "--format",
        "-f",
        choices=["txt", "docx", "json"],
        default="txt",
        help="Output format (default: txt)"
    )
    parser.add_argument(
        "--interactive",
        "-i",
        action="store_true",
        help="Interactive mode"
    )
    
    args = parser.parse_args()
    
    formatter = SummaryFormatter()
    
    # Interactive mode
    if args.interactive or not args.input:
        print("=== Interactive Summary Formatter ===")
        print("\nPaste your raw summary text below.")
        print("Press Ctrl+D (Unix) or Ctrl+Z (Windows) when done:\n")
        
        try:
            raw_text = sys.stdin.read()
        except KeyboardInterrupt:
            print("\nCancelled.")
            return
        
        if not raw_text.strip():
            print("No input provided.")
            return
    else:
        # File mode
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"Error: Input file '{args.input}' not found.")
            return
        
        with open(input_path, 'r', encoding='utf-8') as f:
            raw_text = f.read()
    
    # Parse the summary
    print("\nParsing summary...")
    structured_data = formatter.parse_raw_summary(raw_text)
    
    # Generate output
    if args.format == "json":
        output = json.dumps(structured_data, indent=2)
        
    elif args.format == "docx":
        if not DOCX_AVAILABLE:
            print("Error: python-docx is not installed. Use 'txt' or 'json' format instead.")
            return
        
        doc = formatter.format_docx_output(structured_data)
        output_path = args.output or "summary.docx"
        doc.save(output_path)
        print(f"\nWord document saved to: {output_path}")
        return
        
    else:  # txt
        output = formatter.format_text_output(structured_data)
    
    # Output results
    if args.output and args.format != "docx":
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output)
        print(f"\nFormatted summary saved to: {args.output}")
    else:
        print("\n" + output)


if __name__ == "__main__":
    main()
